import socket
import ssl
from urllib.parse import urlparse, parse_qs
import requests
from bs4 import BeautifulSoup
import re
from art import *
from docx import Document


def is_valid_url(url):
    try:
        result = urlparse(url)
        return all([result.scheme, result.netloc])
    except ValueError:
        print("The URL is not valid.")
        return False


def is_url_reachable(url):
    try:
        response = requests.get(url)
        if response.status_code == 200:
            print("The URL is reachable.")
            return True
        else:
            print("The URL is not reachable.")
            return False
    except requests.ConnectionError:
        print("The URL is not reachable.")
        return False


def check_url_redirection(url):
    try:
        response = requests.head(url, allow_redirects=True)

        if response.status_code == 200:  # If status code is 200, it means no redirection
            print("The URL is not a redirection URL.")
            return None
        elif response.history:  # If there are redirects in history, return the final URL
            print("The URL is a redirection URL.")
            return response.url
        else:  # If neither 200 nor any redirects, it's not a redirection
            print("The URL is not a redirection URL.")
            return None
    except requests.ConnectionError:
        print("Error: Connection Error")
        return None


def verify_url_scheme(url, scheme):
    parsed_url = urlparse(url)
    return parsed_url.scheme.lower() == scheme.lower()


def check_url_path(url, paths):
    parsed_url = urlparse(url)
    url_paths = parsed_url.path.split('/')
    return any(path in url_paths for path in paths)


def extract_query_parameters(url):
    parsed_url = urlparse(url)
    query_params = parse_qs(parsed_url.query)
    return query_params


def extract_fragment_identifier(url):
    parsed_url = urlparse(url)
    return parsed_url.fragment


def check_ssl_cert(url):
    try:
        hostname = url.split("://")[-1].split("/")[0]
        ctx = ssl.create_default_context()
        with ctx.wrap_socket(socket.socket(), server_hostname=hostname) as s:
            s.connect((hostname, 443))
            cert = s.getpeercert()
            print(f"Subject: {cert['subject']}")
            print(f"Issuer: {cert['issuer']}")
            print(f"Expiry Date: {cert['notAfter']}")
            print(f"Serial Number: {cert['serialNumber']}")
            print(f"Public Key: {cert['subjectPublicKeyInfo']}")
    except Exception as e:
        print(f"Error checking SSL certificate: {e}")


def check_malware_phishing(url):
    try:
        response = requests.get(url)
        if response.status_code == 200:
            soup = BeautifulSoup(response.text, 'html.parser')
            scripts = soup.find_all('script')
            for script in scripts:
                if re.search(r'malware|virus|phish|trojan', script.get_text(), re.IGNORECASE):
                    print("\033[101Warning: Malware or phishing content detected.\033[0m")
                    return True
            print("No malware or phishing content detected.")
            return False
        else:
            print(f"Failed to fetch webpage. Status code: {response.status_code}")
            return None
    except Exception as e:
        print("An error occurred:", e)
        return None


phishing_domains = [
    "example-phishing.com",
    "phishing-site.org",
    "malicious.site",
    "login-paypal-security.com",
    "bankofamerica-secure-login.net",
    "appleid-verification-support.info",
    "amazon-rewards-alerts.club",
    "netflix-account-verification.xyz",
    "microsoft-support-helpdesk.online",
    "irs-tax-refund-claim.org",
    "google-drive-login-authentication.site",
    "facebook-security-verification-page.com",
    "instagram-account-recovery-page.net",
]


def check_phishing(url):
    try:
        parsed_url = urlparse(url)
        if parsed_url.scheme not in {"https"}:
            print(
                "\033[101mPhishing Alert: This URL does not use HTTP or HTTPS. It might be a phishing attempt.\033[0m")
        elif not parsed_url.netloc:
            print("\033[101mPhishing Alert: Invalid URL.\033[0m")
        elif parsed_url.netloc in phishing_domains:
            print("\033[101mPhishing Alert: This website is a known phishing domain.\033[0m")
        else:
            print("")
            print("                      _____👍👍👍👍👍👍👍______      ")
            print("              \033[92mSafe Website: This website appears to be safe.\033[0m")
    except Exception as e:
        print("\033[101mError:", str(e), "\033[0m")


# Example usage:
print(text2art("Phisher @ info"))
print('''                        
                               $...Welcome to phisher@info tools.....$
                      Your frontline defense in the battle against cyber deception, 
                      ensuring peace of mind in an ever-evolving digital world.
                      ''')
print("")
print("")
user_input = input('''  
    @ Enter 1 to open tool 
    @ Enter 2 for about the tool

    Please Select and Enter  : ''')
url = ""  # Initialize url variable
if user_input == '1':
    url = str(input(" Enter the url :"))
elif user_input == '2':
    print("")
    print("")
    print("                                -@-#&-*_%_4_$-#-*-#-! ( PHISHER@ INFO ) -@-#&-*_%_4_$-#-*-#-!  ")
    print(''' 


 URL Phishing Detection:
                        This function checks if a website might be trying to trick you into giving away 
                         your sensitive information, like passwords or credit card numbers.It looks at 
                          the website's address to see if it uses a secure connection (HTTPS). If it doesn't, it might be risky.
                           It also checks if the website's address looks suspicious or similar to known phishing websites.
                           If the website seems safe after these checks, it tells you it's safe to visit. Otherwise,
                            it warns you about potential dangers.

 URL Syntax:
              This function checks if a given string is a valid URL by attempting to parse it using 
               the urlparse function from the urllib.parse module.It verifies if both the scheme and
                netloc components are present,indicating a valid URL structure.


 URL Reachability:
                     This function checks if a URL is reachable by making an HTTP GET request to 
                      it using the requests.get() function from the requests module.
                        It returns True if the HTTP response status code is 200 (OK), 
                          indicating that the URL is reachable.

 URL Redirection:
                 This function checks if a URL redirects to another location. 
                   It sends a HEAD request to the URL with allow_redirects=True parameter 
                    to follow redirections,and then checks if there are any redirections in 
                      the response history If there are redirections,it returns 
                        the final redirected URL; otherwise, it returns None.                             

 URL Scheme: 
            This function verifies if the URL is using a specific scheme (e.g., HTTP or HTTPS).
              It parses the URL using urlparse and compares the scheme component 
                to the specified scheme (case-insensitive).                                 

 URL Path:
           This function checks specific paths or components in the URL. 
            It parses the URL and compares each component of the path with the provided list of paths.
             It returns True if any of the specified paths are found in the URL path.                            

 URL Parameters:
                This function parses and extracts query parameters from the URL.
                 It uses parse_qs function from urllib.parse module to parse 
                   the query string component of the URL and returns a dictionary containing 
                    the parameter names as keys and lists of parameter values as values.

 URL Fragment: 
               This function extracts the fragment identifier (part of the URL after #).
                It parses the URL using urlparse and returns the fragment component.


 URL Malware: 
               This function checks if a webpage contains any harmful content like malware or phishing attempts.
                It visits the webpage specified by the provided URL.
                 If the webpage is accessible (status code 200), it scans the webpage's HTML code for suspicious 
                  keywords often associated with malware.If any of these keywords are found, it warns 
                  the user about potential threats. Otherwise, it assures the user that no harmful content was detected              

 URL SSL Certificate:               
                     This function examines the SSL certificate of a website to verify its security.
                      It establishes a secure connection with the website using HTTPS.
                       Once connected, it retrieves and analyzes the SSL certificate details including the issuing authority, 
                        expiration date, and public key.It provides this information to the user, ensuring transparency 
                         about the website's security features.
                         by sam@cm
''')
    input("Press Enter to exit...")
    exit()
else:
    print("Incorrect option. Please enter either 1 or 2.")

print("\n")
print("----CHECKING---PHISHING--URL---")
print("Checking phishing URL:", check_phishing(url))
print("------URL----VALID---CHECK----\n")
if is_valid_url(url):
    print("The URL is valid.")
print("Is Valid URL:", is_valid_url(url))
print("\n")
print("-----URL----REACHABLE---CHECK----\n")
print("Is URL Reachable:", is_url_reachable(url))
print("\n")
print("----URL---REDIRECTION---CHECK-----\n")
print("Redirected URL:", check_url_redirection(url))
print("\n")
print("-----URL---HTTP---HTTPS---CHECK----\n")
print("Is URL using HTTPS:", verify_url_scheme(url, "https"))
print("\n")
print("-----URL---PATH----CHECK------\n")
print("Contains Path:", check_url_path(url, ["path", "to", "page"]))
print("\n")
print("-----URL---QUERY---PARAMETERS----\n")
print("Query Parameters:", extract_query_parameters(url))
print("\n")
print("-----URL---FRAGMENT---IDENTIFIER------\n")
print("Fragment Identifier:", extract_fragment_identifier(url))
print("\n")
print("-----------------malware----info--------------\n")
print("Is checking malware:", check_malware_phishing(url))
print("\n")
print("-------SSL---CERTIFICATE---DETAIL------\n")
print("Is checking ssl cert:", check_ssl_cert(url))

print("---------------------------------------------------------------------------------------------")
print("\n")

# Perform all checks and store results
results = {
    "URL": url,
    "Valid URL": is_valid_url(url),
    "URL Reachable": is_url_reachable(url),
    "Redirected URL": check_url_redirection(url),
    "Using HTTPS": verify_url_scheme(url, "https"),
    "Contains Path": check_url_path(url, ["path", "to", "page"]),
    "Query Parameters": extract_query_parameters(url),
    "Fragment Identifier": extract_fragment_identifier(url),
    "Contains Malware": check_malware_phishing(url),

}

# Get the file path from the user
print("\n")
print("Like ...C:\sam\ ")
file_path = input("Enter the file path to save the Word document: ")

# Save DataFrame to Word file
output_file = file_path + "result.docx"

# Create a new Word document
doc = Document()

# Add a title to the document
doc.add_heading("--------PHISHER@INFO---REPORT----", level=1)
doc.add_heading("URL Check Results", level=2)
# Add results to the document
for column, value in results.items():
    doc.add_paragraph(f"{column}: {value}")

doc.add_paragraph("Thank you for using Phisher@Info tools!")
# Save the document
doc.save(output_file)
print("\n")
print("Results saved to:", output_file)
print("\n ")
input("Press Enter to exit...")
